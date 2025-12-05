from smartdz import models, fields, api
from smartdz.exceptions import ValidationError
import base64
import tempfile

class DwMeetingPvWizard(models.TransientModel):
    _name = 'dw.meeting.pv.wizard'
    _description = 'Meeting PV Export Wizard'

    meeting_id = fields.Many2one('dw.meeting', required=True)
    export_type = fields.Selection([
        ('pdf', 'PDF'),
        ('word', 'Word (DOCX)'),
    ], default='pdf', required=True)

    file = fields.Binary("File", readonly=True)
    filename = fields.Char("Filename")

    def action_download_pv(self):
        self.ensure_one()
        meeting = self.meeting_id

        if not meeting.pv:
            raise ValidationError("PV content is empty.")

        # PDF
        if self.export_type == "pdf":
            pdf_bytes = meeting._generate_pv_pdf()
            file_data = base64.b64encode(pdf_bytes)
            filename = f"PV_{meeting.name}.pdf"

        # WORD
        else:
            try:
                from docx import Document
            except ImportError:
                raise ValidationError("Please install python-docx.")

            doc = Document()
            doc.add_heading(f"PV - {meeting.name}", level=1)
            doc.add_paragraph(meeting.pv or "")

            temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(temp.name)

            with open(temp.name, 'rb') as f:
                file_data = base64.b64encode(f.read())

            filename = f"PV_{meeting.name}.docx"

        self.write({
            'file': file_data,
            'filename': filename,
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f"/web/content/?model=dw.meeting.pv.wizard&id={self.id}&field=file&filename={filename}&download=true",
            'target': 'self',
        }
